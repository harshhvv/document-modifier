import openpyxl
from docx import Document

# Load the Excel file
wb = openpyxl.load_workbook("mapping.xlsx")
sheet = wb.active

# Load doc1 and doc2
doc1 = Document("doc1.docx")
doc2 = Document("doc2.docx")

# Create a mapping dictionary from Excel
mapping = {}
for row in sheet.iter_rows(values_only=True):
    mapping[row[0]] = row[1]


# Function to find and return paragraphs under a given heading
def get_paragraphs(document, heading):
    paragraphs = []
    found_heading = False
    for paragraph in document.paragraphs:
        if paragraph.style.name.startswith("Heading") and paragraph.text == heading:
            found_heading = True
        elif found_heading and paragraph.style.name.startswith("Heading"):
            break
        elif found_heading:
            paragraphs.append(paragraph.text)
    return paragraphs


# Iterate through headings in doc2 and copy paragraphs from doc1 based on mapping
for heading in mapping.keys():
    paragraphs = get_paragraphs(doc1, heading)
    if paragraphs:
        # Find the corresponding heading in doc2
        for i, paragraph in enumerate(doc2.paragraphs):
            if (
                paragraph.style.name.startswith("Heading")
                and paragraph.text == mapping[heading]
            ):
                # Insert paragraphs below the heading in doc2
                for para in paragraphs:
                    doc2.paragraphs[i].insert_paragraph_before(para)
                break

# Save the updated doc2
doc2.save("updated_doc2.docx")
